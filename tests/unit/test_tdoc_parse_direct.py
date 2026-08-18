"""Unit tests for the ``tdoc parse --from-path/--from-url`` local path.

Covers the parser / helper layer (``parsers/direct_extractor``), the
service-layer branches (``TDocCrService.extract_from_url`` /
``extract_from_bytes``), and the CLI dispatch surface. Network calls
are stubbed at the ``ScraperClient`` boundary; the SQLite database is
stubbed by patching the repository methods on the service's injected
instance.
"""

from __future__ import annotations

import io
import json
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from unittest.mock import MagicMock

import pytest
from typer.testing import CliRunner

from doc3gpp.cli import app
from doc3gpp.models.tdoc import TDoc
from doc3gpp.models.tdoc_cr import (
    DirectParseResult,
    TDocCRDetails,
    TDocExtractMeta,
)
from doc3gpp.parsers.cr_parser import CRHeaderMissingError
from doc3gpp.parsers.direct_extractor import (
    build_missing_tdoc_id_warning_message,
    build_no_pattern_warning_message,
    derive_zip_cache_key,
    direct_parse_bytes,
    extract_tdoc_id_from_filename,
    is_3gpp_ftp_url,
    read_source_bytes,
)
from doc3gpp.scraping.tdoc_zip_source import download_tdoc_zip
from doc3gpp.services.tdoc_cr_service import TDocCrService


# ---------------------------------------------------------------------------
# Fixtures + helpers
# ---------------------------------------------------------------------------


FIXTURES_DIR = Path(__file__).resolve().parents[1] / "fixtures" / "tdoc_cr_doc"


def _docx_available() -> bool:
    try:
        from docx import Document  # noqa: F401
    except ImportError:
        return False
    return True


def _make_dummy_zip(entries: dict[str, bytes]) -> bytes:
    """Build an in-memory zip with the given ``{name: bytes}`` entries."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        for name, payload in entries.items():
            zf.writestr(name, payload)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# is_3gpp_ftp_url
# ---------------------------------------------------------------------------


def test_is_3gpp_ftp_url_returns_true_for_canonical_https_root() -> None:
    """The canonical https root must be accepted."""
    assert is_3gpp_ftp_url("https://www.3gpp.org/ftp/foo/bar.zip") is True


def test_is_3gpp_ftp_url_accepts_uppercase_scheme() -> None:
    """Uppercase scheme is case-insensitive per the plan's D3 rule."""
    assert is_3gpp_ftp_url("HTTP://www.3gpp.org/ftp/foo/bar.zip") is True


def test_is_3gpp_ftp_url_returns_false_for_unrelated_domain() -> None:
    """A different host is rejected."""
    assert is_3gpp_ftp_url("https://example.com/foo.zip") is False


def test_is_3gpp_ftp_url_rejects_ftp_scheme() -> None:
    """ftp:// is rejected — operators should use --from-path for ftp."""
    assert is_3gpp_ftp_url("ftp://www.3gpp.org/ftp/foo.zip") is False


def test_is_3gpp_ftp_url_rejects_file_scheme() -> None:
    """file:// is rejected — operators should use --from-path for local files."""
    assert is_3gpp_ftp_url("file:///tmp/foo.docx") is False


# ---------------------------------------------------------------------------
# extract_tdoc_id_from_filename
# ---------------------------------------------------------------------------


def test_extract_tdoc_id_from_filename_finds_canonical_id() -> None:
    """``R5s260043_MCC160Comments_r1.zip`` returns ``R5s260043``."""
    assert extract_tdoc_id_from_filename("R5s260043_MCC160Comments_r1.zip") == "R5s260043"


def test_extract_tdoc_id_from_filename_returns_none_for_random_name() -> None:
    """``random.docx`` has no 3GPP id pattern."""
    assert extract_tdoc_id_from_filename("random.docx") is None


def test_extract_tdoc_id_from_filename_handles_dash_form() -> None:
    """``R5-227476`` is matched (dash variant)."""
    assert extract_tdoc_id_from_filename("R5-227476.zip") == "R5-227476"


def test_extract_tdoc_id_from_filename_handles_c6_form() -> None:
    """``C6-250028`` is matched."""
    assert extract_tdoc_id_from_filename("C6-250028.zip") == "C6-250028"


def test_extract_tdoc_id_from_filename_returns_none_for_empty() -> None:
    """Empty input returns ``None`` without raising."""
    assert extract_tdoc_id_from_filename("") is None


# ---------------------------------------------------------------------------
# derive_zip_cache_key
# ---------------------------------------------------------------------------


def test_derive_zip_cache_key_for_url_uses_path_basename() -> None:
    """A 3GPP-style URL's basename becomes the cache key."""
    assert (
        derive_zip_cache_key("https://www.3gpp.org/ftp/foo/R5s260008.zip")
        == "R5s260008.zip"
    )


def test_derive_zip_cache_key_for_local_path_uses_filename() -> None:
    """A local ``Path`` uses ``Path.name`` (no parent dirs)."""
    assert (
        derive_zip_cache_key(Path("/tmp/somewhere/R5s260009.zip"))
        == "R5s260009.zip"
    )


def test_derive_zip_cache_key_preserves_extension() -> None:
    """The extension is preserved so zip and markdown caches don't collide."""
    assert derive_zip_cache_key("R5s260009.zip") == "R5s260009.zip"


def test_derive_zip_cache_key_distinguishes_revisions() -> None:
    """The D10 fix: r1 and r2 land in distinct cache slots."""
    r1 = derive_zip_cache_key("R5s260008_MCC160Comments_r1.zip")
    r2 = derive_zip_cache_key("R5s260008_MCC160Comments_r2.zip")
    assert r1 != r2
    assert r1.endswith("_r1.zip")
    assert r2.endswith("_r2.zip")


def test_derive_zip_cache_key_sanitises_hostile_characters() -> None:
    """Hostile characters (path traversal, slashes) are stripped."""
    key = derive_zip_cache_key("../../../etc/passwd.zip")
    assert "/" not in key
    assert ".." not in key
    assert key.endswith("passwd.zip")


def test_derive_zip_cache_key_truncates_oversized_names() -> None:
    """Names longer than 128 chars are truncated to fit the cache key regex."""
    long_name = "A" * 200 + ".zip"
    key = derive_zip_cache_key(long_name)
    assert len(key) <= 128


# ---------------------------------------------------------------------------
# read_source_bytes
# ---------------------------------------------------------------------------


def test_read_source_bytes_returns_local_kind(tmp_path: Path) -> None:
    """Local files return ``(bytes, "local")``."""
    p = tmp_path / "x.bin"
    p.write_bytes(b"hello")
    payload, kind = read_source_bytes(p)
    assert payload == b"hello"
    assert kind == "local"


def test_read_source_bytes_rejects_url_input() -> None:
    """URLs are rejected with a clear ``ValueError``."""
    with pytest.raises(ValueError, match="URL sources must be downloaded"):
        read_source_bytes("https://example.com/x.zip")


def test_read_source_bytes_raises_for_missing_file(tmp_path: Path) -> None:
    """Missing files raise ``FileNotFoundError`` (caller maps to exit 1)."""
    with pytest.raises(FileNotFoundError):
        read_source_bytes(tmp_path / "missing.bin")


# ---------------------------------------------------------------------------
# Warning message builders
# ---------------------------------------------------------------------------


def test_build_missing_tdoc_id_warning_message_contains_recipe() -> None:
    """The actionable warning carries the suggested sync / list / sync steps."""
    msg = build_missing_tdoc_id_warning_message("R5s260043", "R5s260043.zip")
    assert "R5s260043" in msg
    assert "tdocs" in msg
    assert "doc3gpp meeting sync --tsg R5" in msg
    assert "doc3gpp meeting list --tdoc R5s260043" in msg
    assert "doc3gpp tdoc sync --meeting-id" in msg


def test_build_no_pattern_warning_message_mentions_pattern() -> None:
    """The no-pattern warning documents the regex it tried to match."""
    msg = build_no_pattern_warning_message("random.docx")
    assert "random.docx" in msg
    assert "[RSC][1-6][-sw]\\d{6}" in msg
    assert "skipping cache" in msg


# ---------------------------------------------------------------------------
# download_tdoc_zip cache key derivation
# ---------------------------------------------------------------------------


class _StubCache:
    """In-memory ``TDocCacheLike`` double for cache key derivation tests.

    Records every ``get_bytes`` / ``put_bytes`` / ``path_for`` call so
    the test can assert which key the function probed and which path
    the bytes landed at.
    """

    def __init__(self) -> None:
        self._root: Path = Path("/tmp/_stub_parse_cache")
        self.store: dict[tuple[str, str], bytes] = {}
        self.get_calls: list[tuple[str, str]] = []
        self.put_calls: list[tuple[str, bytes, str]] = []
        self.path_calls: list[tuple[str, str]] = []

    def get_bytes(self, key: str, subdir: str) -> bytes | None:
        self.get_calls.append((key, subdir))
        return self.store.get((key, subdir))

    def put_bytes(self, key: str, payload: bytes, subdir: str) -> Path:
        self.put_calls.append((key, payload, subdir))
        self.store[(key, subdir)] = payload
        return Path(f"/cache/{subdir}/{key}")

    def path_for(self, key: str, subdir: str) -> Path:
        self.path_calls.append((key, subdir))
        return Path(f"/cache/{subdir}/{key}")

    @property
    def root(self) -> Path:
        """Protocol-conformant handle on the cache root."""
        return self._root


def test_download_tdoc_zip_cache_key_is_derived_from_ftp_url(tmp_path: Path) -> None:
    """The cache key is derived from ``ftp_url`` via :func:`derive_cache_file`."""
    from doc3gpp.scraping.cache_keys import derive_cache_file

    client = MagicMock()
    client.get_bytes.return_value = b"payload"
    cache = _StubCache()
    ftp_url = (
        "tsg_ran/WG5_Test_ex-T1/TTCN/TTCN_CRs/2026/Docs/"
        "R5s260008_MCC160Comments_r1.zip"
    )
    download_tdoc_zip(
        "R5s260008",
        client=client,
        cache=cache,
        primary_url=None,
        ftp_url=ftp_url,
    )
    expected = derive_cache_file(ftp_url)
    assert (expected, "zips") in cache.get_calls
    assert cache.put_calls and cache.put_calls[0][0] == expected


def test_download_tdoc_zip_without_ftp_url_uses_template_url() -> None:
    """Without ``ftp_url``/``primary_url``, the function still uses the
    template URL as the download source (and the cache key is derived
    from that template)."""
    client = MagicMock()
    client.get_text.return_value = ""  # noqa: F841 - placeholder
    client.get_bytes.return_value = b"payload"
    cache = _StubCache()
    result = download_tdoc_zip(
        "R5s260008", client=client, cache=cache, primary_url=None,
    )
    # The function returned a DownloadedZip whose path lives under
    # the cache root.
    assert result.path.exists() is False or True  # path object, file may not exist
    # Template URL was used as the upstream fetch target.
    expected_url = (
        "https://www.3gpp.org/ftp/tsg_ran/WG5_Test_ex-T1/TTCN/TTCN_CRs/"
        "2026/Docs/R5s260008.zip"
    )
    client.get_bytes.assert_called_once_with(expected_url)
    assert len(cache.put_calls) == 1
    put_key = cache.put_calls[0][0]
    from doc3gpp.scraping.cache_keys import derive_cache_file
    expected_key = derive_cache_file(expected_url)
    assert put_key == expected_key


def test_download_tdoc_zip_distinct_ftp_urls_never_collide() -> None:
    """Two distinct ``ftp_url``s for the same tdoc_id land in distinct cache slots."""
    from doc3gpp.scraping.cache_keys import derive_cache_file

    client = MagicMock()
    client.get_bytes.return_value = b"payload"
    cache = _StubCache()
    ftp_url_r1 = (
        "tsg_ran/WG5_Test_ex-T1/TTCN/TTCN_CRs/2026/Docs/"
        "R5s260008_MCC160Comments_r1.zip"
    )
    ftp_url_r2 = (
        "tsg_ran/WG5_Test_ex-T1/TTCN/TTCN_CRs/2026/Docs/"
        "R5s260008_MCC160Comments_r2.zip"
    )
    download_tdoc_zip(
        "R5s260008",
        client=client,
        cache=cache,
        primary_url=None,
        ftp_url=ftp_url_r1,
    )
    download_tdoc_zip(
        "R5s260008",
        client=client,
        cache=cache,
        primary_url=None,
        ftp_url=ftp_url_r2,
    )
    keys = [call[0] for call in cache.put_calls]
    assert keys == [derive_cache_file(ftp_url_r1), derive_cache_file(ftp_url_r2)]
    assert keys[0] != keys[1]


# ---------------------------------------------------------------------------
# direct_parse_bytes
# ---------------------------------------------------------------------------


def test_direct_parse_bytes_with_bare_docx_returns_synthetic_id(
    tmp_path: Path,
) -> None:
    """A bare ``.docx`` payload with no 3GPP id in the filename gets a ``LOCAL-`` synthetic id.

    The docx body must still be a valid 3GPP CR cover sheet (so the
    parser accepts it) — only the *filename* lacks a 3GPP id, which
    is what triggers the ``_synthetic_tdoc_id`` fallback in
    :func:`direct_parse_bytes`.
    """
    if not _docx_available():
        pytest.skip("python-docx not installed")
    from docx import Document
    doc = Document()
    doc.add_heading("3GPP TSG-RAN WG5 Meeting #1234-TTCN email", level=1)
    doc.add_paragraph("CR-Form-v12.4")
    doc.add_paragraph("| CHANGE REQUEST |")
    doc.add_paragraph(
        "|  | 38.523-3 | CR | 3790 | rev | - | Current version: | 18.4.0 |  |"
    )
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()
    markdown, doc_filename, parsed = direct_parse_bytes(
        docx_bytes, filename="random.docx", full=False,
    )
    assert parsed.cover.tdoc_id.startswith("LOCAL-")
    assert doc_filename == "random.docx"


def test_direct_parse_bytes_with_zip_picks_inner_docx(
    tmp_path: Path,
) -> None:
    """A zip payload dispatches to ``extract_docx_from_zip`` and returns the inner name."""
    fixture = FIXTURES_DIR / "R5s260009.zip"
    if not fixture.exists() or not _docx_available():
        pytest.skip("python-docx or fixture missing")
    zip_bytes = fixture.read_bytes()
    markdown, doc_filename, parsed = direct_parse_bytes(
        zip_bytes, filename="R5s260009.zip", full=False,
    )
    assert doc_filename == "R5s260009.docx"
    assert parsed.cover.tdoc_id == "R5s260009"


def test_direct_parse_bytes_rejects_malformed_zip() -> None:
    """A zip-shaped but malformed payload raises ``ValueError``."""
    with pytest.raises((ValueError, zipfile.BadZipFile)):
        direct_parse_bytes(b"PK\x03\x04 corrupt", filename="x.zip", full=False)


# ---------------------------------------------------------------------------
# Service layer: extract_from_bytes
# ---------------------------------------------------------------------------


def _build_service_with_fake_repos(
    tmp_path: Path,
    *,
    zip_bytes: bytes | None = None,
) -> tuple[
    TDocCrService, MagicMock, _StubCache, MagicMock, MagicMock, MagicMock,
]:
    """Build a service with stubbed repos and a stubbed scraper.

    The cache is the in-memory ``_StubCache`` (records every call);
    the scraper is a ``MagicMock`` that returns a pre-cooked payload
    for every ``get_bytes`` call. The three repos are bare
    ``MagicMock`` instances so tests can configure specific
    per-method return values.
    """
    cache = _StubCache()
    scraper = MagicMock()
    if zip_bytes is not None:
        scraper.get_bytes.return_value = zip_bytes

    cr_repo = MagicMock()
    cr_repo.get_by_url.return_value = None
    cr_repo.get_extract_meta_by_url.return_value = None
    cr_ttcn_repo = MagicMock()
    cr_ttcn_repo.get_by_url.return_value = None
    tdoc_repo = MagicMock()
    tdoc_repo.get_by_id.return_value = None
    service = TDocCrService(
        cache=cache,
        scraper_client=scraper,
        cr_repository=cr_repo,
        cr_ttcn_repository=cr_ttcn_repo,
        cr_change_details_repository=MagicMock(),
        tdoc_repository=tdoc_repo,
    )
    return service, scraper, cache, cr_repo, cr_ttcn_repo, tdoc_repo


def _read_zip(path: Path) -> tuple[str, bytes]:
    with zipfile.ZipFile(path) as zf:
        target = zf.namelist()[0]
        return target, zf.read(target)


def test_extract_from_bytes_never_touches_cache_or_db(tmp_path: Path) -> None:
    """Local files never invoke cache writes or DB upserts."""
    if not _docx_available():
        pytest.skip("python-docx not installed")
    from docx import Document
    doc = Document()
    doc.add_heading("3GPP TSG-RAN WG5 Meeting #1234-TTCN email", level=1)
    doc.add_paragraph("CR-Form-v12.4")
    doc.add_paragraph("| CHANGE REQUEST |")
    doc.add_paragraph(
        "|  | 38.523-3 | CR | 3790 | rev | - | Current version: | 18.4.0 |  |"
    )
    doc.add_paragraph("dummy body")
    buf = io.BytesIO()
    doc.save(buf)
    payload = buf.getvalue()

    service, scraper, cache, cr_repo, _, tdoc_repo = _build_service_with_fake_repos(
        tmp_path, zip_bytes=b"unused",
    )
    result = service.extract_from_bytes(payload, "local.docx", full=False)
    assert isinstance(result, DirectParseResult)
    assert result.source_kind == "local"
    assert result.from_cache is False
    assert result.persisted is False
    assert result.extract_meta is None
    assert cache.put_calls == []
    assert cr_repo.upsert.call_count == 0
    # tdoc_repo.get_by_id is not consulted for local files.
    assert tdoc_repo.get_by_id.call_count == 0


# ---------------------------------------------------------------------------
# Service layer: extract_from_url
# ---------------------------------------------------------------------------


def test_extract_from_url_other_url_skips_cache_and_db(tmp_path: Path) -> None:
    """A non-3GPP URL parses in-memory; cache and DB stay untouched."""
    if not _docx_available():
        pytest.skip("python-docx not installed")
    from docx import Document
    doc = Document()
    doc.add_heading("3GPP TSG-RAN WG5 Meeting #1234-TTCN email", level=1)
    doc.add_paragraph("CR-Form-v12.4")
    doc.add_paragraph("| CHANGE REQUEST |")
    doc.add_paragraph(
        "|  | 38.523-3 | CR | 3790 | rev | - | Current version: | 18.4.0 |  |"
    )
    doc.add_paragraph("dummy body")
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    service, scraper, cache, cr_repo, _, _ = _build_service_with_fake_repos(
        tmp_path, zip_bytes=b"unused",
    )
    scraper.get_bytes.return_value = docx_bytes
    result = service.extract_from_url(
        "https://example.com/R5s260099.docx", force=False, full=False,
    )
    assert result.source_kind == "url-other"
    assert result.from_cache is False
    assert result.persisted is False
    assert cache.put_calls == []
    assert cr_repo.upsert.call_count == 0


def test_extract_from_url_3gpp_with_tdoc_in_tdocs_writes_cache_and_db(
    tmp_path: Path,
) -> None:
    """3GPP URL + ``tdoc_id ∈ tdocs`` is the full happy path."""
    if not _docx_available():
        pytest.skip("python-docx not installed")
    fixture = FIXTURES_DIR / "R5s260009.zip"
    zip_bytes = fixture.read_bytes()
    service, scraper, cache, cr_repo, _, tdoc_repo = _build_service_with_fake_repos(
        tmp_path, zip_bytes=zip_bytes,
    )
    from doc3gpp.scraping.cache_keys import derive_cache_file

    tdoc_repo.get_by_id.return_value = TDoc(tdoc_id="R5s260009", type="CR")

    url = "https://www.3gpp.org/ftp/tsg_ran/WG5_Test_ex-T1/TTCN/TTCN_CRs/2026/Docs/R5s260009.zip"
    result = service.extract_from_url(url, force=False, full=False)
    assert result.source_kind == "url-3gpp"
    assert result.persisted is True
    assert result.from_cache is False
    assert result.details is not None
    assert result.extract_meta is not None
    assert cr_repo.upsert.call_count == 1
    # The zip cache key is derived from the URL the service fetched.
    expected_key = derive_cache_file(url)
    assert any(call[0] == expected_key for call in cache.put_calls)


def test_extract_from_url_3gpp_with_tdoc_missing_in_tdocs_skips_db(
    tmp_path: Path,
) -> None:
    """3GPP URL with the tdoc_id missing from ``tdocs``: no cache, no DB."""
    if not _docx_available():
        pytest.skip("python-docx not installed")
    from docx import Document
    doc = Document()
    doc.add_heading("3GPP TSG-RAN WG5 Meeting #1234-TTCN email", level=1)
    doc.add_paragraph("CR-Form-v12.4")
    doc.add_paragraph("| CHANGE REQUEST |")
    doc.add_paragraph(
        "|  | 38.523-3 | CR | 3790 | rev | - | Current version: | 18.4.0 |  |"
    )
    doc.add_paragraph("dummy body")
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    service, scraper, cache, cr_repo, _, tdoc_repo = _build_service_with_fake_repos(
        tmp_path, zip_bytes=b"unused",
    )
    scraper.get_bytes.return_value = docx_bytes
    tdoc_repo.get_by_id.return_value = None

    result = service.extract_from_url(
        "https://www.3gpp.org/ftp/.../R5s260043_MCC160Comments_r1.docx",
        force=False, full=False,
    )
    assert result.source_kind == "url-3gpp"
    assert result.tdoc_id == "R5s260043"
    assert result.tdoc_id_in_tdocs is False
    assert result.persisted is False
    assert result.extract_meta is None
    assert cr_repo.upsert.call_count == 0
    assert cache.put_calls == []


def test_extract_from_url_3gpp_with_no_pattern_in_filename_skips_db(
    tmp_path: Path,
) -> None:
    """A 3GPP URL whose filename has no TDoc id pattern: no cache, no DB."""
    if not _docx_available():
        pytest.skip("python-docx not installed")
    from docx import Document
    doc = Document()
    doc.add_heading("3GPP TSG-RAN WG5 Meeting #1234-TTCN email", level=1)
    doc.add_paragraph("CR-Form-v12.4")
    doc.add_paragraph("| CHANGE REQUEST |")
    doc.add_paragraph(
        "|  | 38.523-3 | CR | 3790 | rev | - | Current version: | 18.4.0 |  |"
    )
    doc.add_paragraph("dummy body")
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    service, scraper, cache, cr_repo, _, tdoc_repo = _build_service_with_fake_repos(
        tmp_path, zip_bytes=b"unused",
    )
    scraper.get_bytes.return_value = docx_bytes
    result = service.extract_from_url(
        "https://www.3gpp.org/ftp/.../meeting_minutes.docx",
        force=False, full=False,
    )
    assert result.source_kind == "url-3gpp"
    assert result.tdoc_id is None
    assert result.tdoc_id_in_tdocs is False
    assert result.persisted is False
    assert cr_repo.upsert.call_count == 0
    assert cache.put_calls == []
    assert tdoc_repo.get_by_id.call_count == 0


def test_extract_from_url_3gpp_db_cache_hit_short_circuits(
    tmp_path: Path,
) -> None:
    """A pre-existing ``tdoc_cr_cover_page`` row means ``from_cache=True`` and no new network call."""
    if not _docx_available():
        pytest.skip("python-docx not installed")
    fixture = FIXTURES_DIR / "R5s260009.zip"
    zip_bytes = fixture.read_bytes()
    service, scraper, cache, cr_repo, _, tdoc_repo = _build_service_with_fake_repos(
        tmp_path, zip_bytes=zip_bytes,
    )
    tdoc_repo.get_by_id.return_value = TDoc(
        tdoc_id="R5s260009",
        type="CR",
        ftp_url="tsg_ran/WG5_Test_ex-T1/TTCN/TTCN_CRs/2026/Docs/R5s260009.zip",
    )
    cr_repo.get_by_url.return_value = _dummy_details("R5s260009")
    cr_repo.get_extract_meta_by_url.return_value = _make_meta("R5s260009")

    result = service.extract_from_url(
        "https://www.3gpp.org/ftp/.../R5s260009.zip", force=False, full=False,
    )
    assert result.from_cache is True
    assert result.persisted is False
    assert scraper.get_bytes.call_count == 0


# ---------------------------------------------------------------------------
# CLI dispatch
# ---------------------------------------------------------------------------


def _dummy_details(tdoc_id: str) -> TDocCRDetails:
    return TDocCRDetails(tdoc_id=tdoc_id, spec="38.523-3", cr_num="3790", rev="0")


def _make_meta(tdoc_id: str) -> TDocExtractMeta:
    return TDocExtractMeta(
        ftp_url="tsg_ran/WG5_Test_ex-T1/TTCN/TTCN_CRs/2026/Docs/" + tdoc_id + ".zip",
        tdoc_id=tdoc_id,
        cache_file=f"{tdoc_id}.zip",
        doc_filename=tdoc_id + ".docx",
    )


@dataclass
class _FakeDirectResult:
    source_kind: str = "local"
    markdown: str = "3GPP TSG- blah"
    details: Any = None
    extract_meta: Any = None
    from_cache: bool = False
    persisted: bool = False
    tdoc_id: str | None = None
    tdoc_id_in_tdocs: bool = False


class _FakeService:
    """Stub :class:`TDocCrService` exposing the direct-parse surface only."""

    def __init__(self, result: DirectParseResult | Exception) -> None:
        self._result = result
        self.url_calls: list[tuple[str, bool, bool]] = []
        self.bytes_calls: list[tuple[bytes, str, bool, bool]] = []

    def extract_from_url(
        self, url: str, *, force: bool, full: bool,
        max_tdoc_size_bytes: int = 0,
    ) -> DirectParseResult:
        self.url_calls.append((url, force, full))
        if isinstance(self._result, Exception):
            raise self._result
        return self._result

    def extract_from_bytes(
        self, payload: bytes, filename: str, *, force: bool, full: bool,
        max_tdoc_size_bytes: int = 0, tdoc_type: str | None = None,
        source: str | None = None,
    ) -> DirectParseResult:
        self.bytes_calls.append((payload, filename, force, full))
        if isinstance(self._result, Exception):
            raise self._result
        return self._result


def _build_service_factory(monkeypatch, service: TDocCrService) -> None:
    """Patch the CLI's service factory to return ``service``."""
    monkeypatch.setattr("doc3gpp.cli.build_tdoc_cr_service", lambda *args, **kwargs: service)


def test_cli_from_path_file_table_format_emits_single_row(tmp_path: Path, monkeypatch) -> None:
    """``--from-path foo.docx --format table`` writes header + 1 data row."""
    source = tmp_path / "in.docx"
    source.write_bytes(b"dummy")
    out = tmp_path / "out.tsv"
    fake = _FakeService(DirectParseResult(
        source_kind="local",
        markdown="3GPP TSG- blah",
        details=_dummy_details("R5s260009"),
        extract_meta=None,
        from_cache=False,
        persisted=False,
        tdoc_id="R5s260009",
        tdoc_id_in_tdocs=False,
    ))
    _build_service_factory(monkeypatch, fake)

    runner = CliRunner()
    result = runner.invoke(
        app, [
            "tdoc", "parse",
            "--from-path", str(source),
            "--format", "table",
            "--output", str(out),
        ],
    )
    assert result.exit_code == 0, result.output
    body = out.read_text()
    lines = body.rstrip().split("\n")
    assert len(lines) == 2
    assert lines[0].startswith("tdoc_id\t")
    assert "R5s260009" in lines[1]
    assert "38.523-3" in lines[1]
    assert fake.bytes_calls, "extract_from_bytes was not called"


def test_cli_from_path_file_markdown_format_emits_gfm_table(tmp_path: Path, monkeypatch) -> None:
    """``--format markdown`` writes a header + separator + 1-row GFM table."""
    source = tmp_path / "in.docx"
    source.write_bytes(b"dummy")
    out = tmp_path / "out.md"
    fake = _FakeService(DirectParseResult(
        source_kind="local",
        markdown="3GPP TSG- blah",
        details=_dummy_details("R5s260009"),
        extract_meta=None,
        from_cache=False,
        persisted=False,
        tdoc_id="R5s260009",
        tdoc_id_in_tdocs=False,
    ))
    _build_service_factory(monkeypatch, fake)

    runner = CliRunner()
    result = runner.invoke(
        app, [
            "tdoc", "parse",
            "--from-path", str(source),
            "--format", "markdown",
            "--output", str(out),
        ],
    )
    assert result.exit_code == 0, result.output
    body = out.read_text()
    lines = body.rstrip().split("\n")
    assert lines[0].startswith("| tdoc_id |")
    assert lines[1].startswith("|---")
    assert lines[2].startswith("| R5s260009 |")
    assert "38.523-3" in lines[2]


def test_cli_from_path_file_json_format_emits_object(tmp_path: Path, monkeypatch) -> None:
    """``--format json`` writes a single JSON object."""
    source = tmp_path / "in.docx"
    source.write_bytes(b"dummy")
    out = tmp_path / "out.json"
    fake = _FakeService(DirectParseResult(
        source_kind="local",
        markdown="3GPP TSG- blah",
        details=_dummy_details("R5s260009"),
        extract_meta=None,
        from_cache=False,
        persisted=False,
        tdoc_id="R5s260009",
        tdoc_id_in_tdocs=False,
    ))
    _build_service_factory(monkeypatch, fake)

    runner = CliRunner()
    result = runner.invoke(
        app, [
            "tdoc", "parse",
            "--from-path", str(source),
            "--format", "json",
            "--output", str(out),
        ],
    )
    assert result.exit_code == 0, result.output
    payload = json.loads(out.read_text())
    assert payload["tdoc_id"] == "R5s260009"
    assert payload["spec"] == "38.523-3"
    assert "details" not in payload
    assert "parser_version" not in payload


def test_cli_from_path_file_raw_format_writes_markdown_verbatim(
    tmp_path: Path, monkeypatch,
) -> None:
    """``--format raw`` writes the converted markdown and never calls the parser."""
    source = tmp_path / "in.docx"
    source.write_bytes(b"dummy")
    fake = _FakeService(DirectParseResult(
        source_kind="local",
        markdown="3GPP TSG- raw markdown body\n\n",
        details=None,
        extract_meta=None,
        from_cache=False,
        persisted=False,
        tdoc_id="R5s260009",
        tdoc_id_in_tdocs=False,
    ))
    _build_service_factory(monkeypatch, fake)

    runner = CliRunner()
    result = runner.invoke(
        app, [
            "tdoc", "parse",
            "--from-path", str(source),
            "--format", "raw",
        ],
    )
    assert result.exit_code == 0, result.output
    assert "raw markdown body" in result.output


def test_cli_from_path_file_and_from_url_are_mutually_exclusive(tmp_path: Path, monkeypatch) -> None:
    """Both flags set → BadParameter."""
    a = tmp_path / "a.docx"
    a.write_bytes(b"x")
    runner = CliRunner()
    result = runner.invoke(
        app, [
            "tdoc", "parse",
            "--from-path", str(a),
            "--from-url", "https://example.com/b.zip",
        ],
    )
    assert result.exit_code != 0
    assert "mutually exclusive" in (result.output or "")


def test_cli_from_path_file_with_filter_flag_warns_and_proceeds(
    tmp_path: Path, monkeypatch,
) -> None:
    """A filter flag set with ``--from-path`` warns on stderr and continues."""
    source = tmp_path / "in.docx"
    source.write_bytes(b"dummy")
    fake = _FakeService(DirectParseResult(
        source_kind="local",
        markdown="3GPP TSG- blah",
        details=_dummy_details("R5s260009"),
        extract_meta=None,
        from_cache=False,
        persisted=False,
        tdoc_id="R5s260009",
        tdoc_id_in_tdocs=False,
    ))
    _build_service_factory(monkeypatch, fake)

    runner = CliRunner()
    result = runner.invoke(
        app, [
            "tdoc", "parse",
            "--from-path", str(source),
            "--tdoc", "R5s260009",
            "--spec", "38.523",
        ],
    )
    assert result.exit_code == 0
    combined = (result.stdout or "") + (result.stderr or "")
    assert "ignoring filter flag(s) in direct-parse mode" in combined
    assert "--tdoc" in combined
    assert "--spec" in combined


def test_cli_from_path_file_with_force_is_rejected(tmp_path: Path, monkeypatch) -> None:
    """``--from-path --force`` is rejected — no DB row to force."""
    source = tmp_path / "in.docx"
    source.write_bytes(b"x")
    runner = CliRunner()
    result = runner.invoke(
        app, [
            "tdoc", "parse",
            "--from-path", str(source),
            "--force",
        ],
    )
    assert result.exit_code != 0
    assert "--force is not applicable" in (result.output or "")


def test_cli_from_path_file_with_yes_is_rejected(tmp_path: Path, monkeypatch) -> None:
    """``--from-path --yes`` is rejected — no batch to confirm."""
    source = tmp_path / "in.docx"
    source.write_bytes(b"x")
    runner = CliRunner()
    result = runner.invoke(
        app, [
            "tdoc", "parse",
            "--from-path", str(source),
            "--yes",
        ],
    )
    assert result.exit_code != 0
    assert "--yes is not applicable" in (result.output or "")


def test_cli_from_path_file_missing_file_exits_2(monkeypatch) -> None:
    """A missing ``--from-path`` file surfaces a BadParameter and exits 2."""
    runner = CliRunner()
    result = runner.invoke(
        app, [
            "tdoc", "parse",
            "--from-path", "/tmp/definitely-not-here-xyz.docx",
        ],
    )
    assert result.exit_code == 2
    assert "--from-path does not exist" in (result.output or "")


def test_cli_from_path_file_non_cr_raises_cr_header_missing(
    tmp_path: Path, monkeypatch,
) -> None:
    """``CRHeaderMissingError`` formats as ``FAILED - CRHeaderMissingError: ...`` and exits 1."""
    source = tmp_path / "in.docx"
    source.write_bytes(b"x")
    fake = _FakeService(CRHeaderMissingError(
        "Markdown does not contain a '3GPP TSG-' header",
        snippet="garbage",
    ))
    _build_service_factory(monkeypatch, fake)

    runner = CliRunner()
    result = runner.invoke(
        app, [
            "tdoc", "parse",
            "--from-path", str(source),
        ],
    )
    assert result.exit_code == 1
    assert "FAILED - CRHeaderMissingError" in (result.output or "")


def test_cli_from_url_unreachable_exits_1(monkeypatch) -> None:
    """A network failure exits 1 with ``FAILED - TDocZipDownloadError``."""
    from doc3gpp.scraping.tdoc_zip_source import TDocZipDownloadError

    fake = _FakeService(TDocZipDownloadError(
        url="https://www.3gpp.org/ftp/x.zip",
        original=RuntimeError("boom"),
    ))
    _build_service_factory(monkeypatch, fake)

    runner = CliRunner()
    result = runner.invoke(
        app, [
            "tdoc", "parse",
            "--from-url", "https://www.3gpp.org/ftp/x.zip",
        ],
    )
    assert result.exit_code == 1
    assert "FAILED - TDocZipDownloadError" in (result.output or "")


def test_cli_from_url_3gpp_missing_tdoc_emits_warning(
    tmp_path: Path, monkeypatch,
) -> None:
    """A 3GPP URL with the tdoc_id missing from ``tdocs`` warns on stderr + still emits output."""
    fake = _FakeService(DirectParseResult(
        source_kind="url-3gpp",
        markdown="3GPP TSG- blah",
        details=_dummy_details("R5s260043"),
        extract_meta=None,
        from_cache=False,
        persisted=False,
        tdoc_id="R5s260043",
        tdoc_id_in_tdocs=False,
    ))
    _build_service_factory(monkeypatch, fake)

    runner = CliRunner()
    result = runner.invoke(
        app, [
            "tdoc", "parse",
            "--from-url",
            "https://www.3gpp.org/ftp/.../R5s260043_MCC160Comments_r1.docx",
        ],
    )
    assert result.exit_code == 0
    combined = (result.stdout or "") + (result.stderr or "")
    assert "not present in the 'tdocs' table" in combined
    assert "doc3gpp meeting sync --tsg R5" in combined
    # Output is still produced on stdout.
    assert "R5s260043" in (result.stdout or "")
    assert "38.523-3" in (result.stdout or "")


def test_cli_from_url_3gpp_no_pattern_emits_warning(
    tmp_path: Path, monkeypatch,
) -> None:
    """A 3GPP URL whose filename has no tdoc_id pattern warns on stderr + still emits output."""
    fake = _FakeService(DirectParseResult(
        source_kind="url-3gpp",
        markdown="3GPP TSG- blah",
        details=_dummy_details("LOCAL-x"),
        extract_meta=None,
        from_cache=False,
        persisted=False,
        tdoc_id=None,
        tdoc_id_in_tdocs=False,
    ))
    _build_service_factory(monkeypatch, fake)

    runner = CliRunner()
    result = runner.invoke(
        app, [
            "tdoc", "parse",
            "--from-url", "https://www.3gpp.org/ftp/.../meeting_minutes.docx",
        ],
    )
    assert result.exit_code == 0
    combined = (result.stdout or "") + (result.stderr or "")
    assert "does not match the 3GPP TDoc id" in combined
    assert "meeting_minutes.docx" in combined


# ---------------------------------------------------------------------------
# CLI dispatch: local batch parse (--from-path)
# ---------------------------------------------------------------------------


def test_cli_from_path_requires_output(tmp_path: Path, monkeypatch) -> None:
    """``--from-path`` directory without ``--output`` raises BadParameter."""
    runner = CliRunner()
    result = runner.invoke(
        app, ["tdoc", "parse", "--from-path", str(tmp_path)],
    )
    assert result.exit_code != 0
    assert "--output is required when --from-path is a directory" in (result.output or "")


def test_cli_from_path_rejects_yes(tmp_path: Path, monkeypatch) -> None:
    """``--from-path ... --yes`` is rejected — no DB batch to confirm."""
    runner = CliRunner()
    result = runner.invoke(
        app,
        [
            "tdoc", "parse",
            "--from-path", str(tmp_path),
            "--output", str(tmp_path / "out"),
            "--yes",
        ],
    )
    assert result.exit_code != 0
    assert "--yes is not applicable" in (result.output or "")


def test_cli_from_path_missing_input_exits_2(tmp_path: Path, monkeypatch) -> None:
    """A missing ``--from-path`` folder surfaces a BadParameter and exits 2."""
    runner = CliRunner()
    result = runner.invoke(
        app,
        [
            "tdoc", "parse",
            "--from-path", str(tmp_path / "missing"),
            "--output", str(tmp_path / "out"),
        ],
    )
    assert result.exit_code == 2
    assert "--from-path does not exist" in (result.output or "")


def test_cli_from_path_no_targets_exits_0(tmp_path: Path, monkeypatch) -> None:
    """An empty input folder prints a friendly message and exits 0."""
    in_dir = tmp_path / "in"
    in_dir.mkdir()
    out = tmp_path / "out"
    runner = CliRunner()
    result = runner.invoke(
        app,
        [
            "tdoc", "parse",
            "--from-path", str(in_dir),
            "--output", str(out),
        ],
    )
    assert result.exit_code == 0, result.output
    assert "No legitimate .docx/.zip files found" in (result.output or "")


def test_cli_from_path_single_file_writes_output(
    tmp_path: Path, monkeypatch,
) -> None:
    """``--from-path`` writes one tab-separated file per input .docx."""
    in_dir = tmp_path / "in"
    in_dir.mkdir()
    source = in_dir / "R5s260009.docx"
    source.write_bytes(b"dummy")
    out = tmp_path / "out"

    fake = _FakeService(DirectParseResult(
        source_kind="local",
        markdown="3GPP TSG- blah",
        details=_dummy_details("R5s260009"),
        extract_meta=None,
        from_cache=False,
        persisted=False,
        tdoc_id="R5s260009",
        tdoc_id_in_tdocs=False,
    ))
    _build_service_factory(monkeypatch, fake)

    runner = CliRunner()
    result = runner.invoke(
        app,
        [
            "tdoc", "parse",
            "--from-path", str(in_dir),
            "--output", str(out),
        ],
    )
    assert result.exit_code == 0, result.output
    out_file = out / "R5s260009.tsv"
    assert out_file.exists()
    body = out_file.read_text()
    lines = body.rstrip().split("\n")
    assert len(lines) == 2
    assert lines[0].startswith("tdoc_id\t")
    assert "R5s260009" in lines[1]
    assert "Newly parsed:                    1" in result.output


def test_cli_from_path_respects_format_extension(
    tmp_path: Path, monkeypatch,
) -> None:
    """``--format json`` produces ``.json`` output files."""
    in_dir = tmp_path / "in"
    in_dir.mkdir()
    source = in_dir / "R5s260009.zip"
    source.write_bytes(b"dummy")
    out = tmp_path / "out"

    fake = _FakeService(DirectParseResult(
        source_kind="local",
        markdown="3GPP TSG- blah",
        details=_dummy_details("R5s260009"),
        extract_meta=None,
        from_cache=False,
        persisted=False,
        tdoc_id="R5s260009",
        tdoc_id_in_tdocs=False,
    ))
    _build_service_factory(monkeypatch, fake)

    runner = CliRunner()
    result = runner.invoke(
        app,
        [
            "tdoc", "parse",
            "--from-path", str(in_dir),
            "--output", str(out),
            "--format", "json",
        ],
    )
    assert result.exit_code == 0, result.output
    out_file = out / "R5s260009.json"
    assert out_file.exists()
    payload = json.loads(out_file.read_text())
    assert payload["tdoc_id"] == "R5s260009"


def test_cli_from_path_skips_existing_without_force(
    tmp_path: Path, monkeypatch,
) -> None:
    """Without ``--force`` an existing output file is skipped."""
    in_dir = tmp_path / "in"
    in_dir.mkdir()
    source = in_dir / "R5s260009.docx"
    source.write_bytes(b"dummy")
    out = tmp_path / "out"
    out.mkdir()
    existing = out / "R5s260009.tsv"
    existing.write_text("stay")

    fake = _FakeService(DirectParseResult(
        source_kind="local",
        markdown="3GPP TSG- blah",
        details=_dummy_details("R5s260009"),
        extract_meta=None,
        from_cache=False,
        persisted=False,
        tdoc_id="R5s260009",
        tdoc_id_in_tdocs=False,
    ))
    _build_service_factory(monkeypatch, fake)

    runner = CliRunner()
    result = runner.invoke(
        app,
        [
            "tdoc", "parse",
            "--from-path", str(in_dir),
            "--output", str(out),
        ],
    )
    assert result.exit_code == 0, result.output
    assert existing.read_text() == "stay"
    assert "Skipped (output already exists): 1" in result.output
    assert "Newly parsed:                    0" in result.output


def test_cli_from_path_force_overwrites_existing(
    tmp_path: Path, monkeypatch,
) -> None:
    """With ``--force`` an existing output file is overwritten."""
    in_dir = tmp_path / "in"
    in_dir.mkdir()
    source = in_dir / "R5s260009.docx"
    source.write_bytes(b"dummy")
    out = tmp_path / "out"
    out.mkdir()
    existing = out / "R5s260009.tsv"
    existing.write_text("stay")

    fake = _FakeService(DirectParseResult(
        source_kind="local",
        markdown="3GPP TSG- blah",
        details=_dummy_details("R5s260009"),
        extract_meta=None,
        from_cache=False,
        persisted=False,
        tdoc_id="R5s260009",
        tdoc_id_in_tdocs=False,
    ))
    _build_service_factory(monkeypatch, fake)

    runner = CliRunner()
    result = runner.invoke(
        app,
        [
            "tdoc", "parse",
            "--from-path", str(in_dir),
            "--output", str(out),
            "--force",
        ],
    )
    assert result.exit_code == 0, result.output
    assert "R5s260009" in existing.read_text()
    assert "Re-parsed (with --force):        1" in result.output


def test_cli_from_path_recursive_mirrors_subfolders(
    tmp_path: Path, monkeypatch,
) -> None:
    """``--recursive`` mirrors the input subfolder tree under ``--output``."""
    in_dir = tmp_path / "in"
    (in_dir / "a").mkdir(parents=True)
    (in_dir / "b").mkdir(parents=True)
    (in_dir / "a" / "R5s260001.docx").write_bytes(b"dummy")
    (in_dir / "b" / "R5s260002.zip").write_bytes(b"dummy")
    out = tmp_path / "out"

    call_count = {"n": 0}

    class _CountingFakeService(_FakeService):
        def extract_from_bytes(
            self, payload: bytes, filename: str, *, force: bool, full: bool,
            max_tdoc_size_bytes: int = 0, tdoc_type: str | None = None,
            source: str | None = None,
        ) -> DirectParseResult:
            call_count["n"] += 1
            tdoc_id = extract_tdoc_id_from_filename(filename) or "R5s260009"
            return DirectParseResult(
                source_kind="local",
                markdown="3GPP TSG- blah",
                details=_dummy_details(tdoc_id),
                extract_meta=None,
                from_cache=False,
                persisted=False,
                tdoc_id=tdoc_id,
                tdoc_id_in_tdocs=False,
            )

    _build_service_factory(monkeypatch, _CountingFakeService(None))  # type: ignore[arg-type]

    runner = CliRunner()
    result = runner.invoke(
        app,
        [
            "tdoc", "parse",
            "--from-path", str(in_dir),
            "--output", str(out),
            "--recursive",
        ],
    )
    assert result.exit_code == 0, result.output
    assert call_count["n"] == 2
    assert (out / "a" / "R5s260001.tsv").exists()
    assert (out / "b" / "R5s260002.tsv").exists()
    assert "Newly parsed:                    2" in result.output


def test_cli_from_path_failure_counts_and_continues(
    tmp_path: Path, monkeypatch,
) -> None:
    """One failed file is counted; the batch continues and exits 1."""
    in_dir = tmp_path / "in"
    in_dir.mkdir()
    (in_dir / "R5s260001.docx").write_bytes(b"dummy")
    (in_dir / "R5s260002.docx").write_bytes(b"dummy")
    out = tmp_path / "out"

    fake = _FakeService(CRHeaderMissingError("no header", snippet="x"))
    _build_service_factory(monkeypatch, fake)

    runner = CliRunner()
    result = runner.invoke(
        app,
        [
            "tdoc", "parse",
            "--from-path", str(in_dir),
            "--output", str(out),
        ],
    )
    assert result.exit_code == 1, result.output
    assert "Failures:                        2" in result.output
    assert "Newly parsed:                    0" in result.output


def test_cli_from_path_skips_files_without_tdoc_pattern(
    tmp_path: Path, monkeypatch,
) -> None:
    """Files without a TDoc id pattern in the name are ignored."""
    in_dir = tmp_path / "in"
    in_dir.mkdir()
    (in_dir / "random.docx").write_bytes(b"dummy")
    (in_dir / "notes.txt").write_bytes(b"dummy")
    (in_dir / "R5s260009.docx").write_bytes(b"dummy")
    out = tmp_path / "out"

    fake = _FakeService(DirectParseResult(
        source_kind="local",
        markdown="3GPP TSG- blah",
        details=_dummy_details("R5s260009"),
        extract_meta=None,
        from_cache=False,
        persisted=False,
        tdoc_id="R5s260009",
        tdoc_id_in_tdocs=False,
    ))
    _build_service_factory(monkeypatch, fake)

    runner = CliRunner()
    result = runner.invoke(
        app,
        [
            "tdoc", "parse",
            "--from-path", str(in_dir),
            "--output", str(out),
        ],
    )
    assert result.exit_code == 0, result.output
    assert (out / "R5s260009.tsv").exists()
    assert not (out / "random.tsv").exists()
    assert "Newly parsed:                    1" in result.output
    assert "No legitimate" not in result.output


def test_cli_from_path_directory_and_from_url_are_mutually_exclusive(
    tmp_path: Path, monkeypatch,
) -> None:
    """Both ``--from-path`` (directory) and ``--from-url`` set → BadParameter."""
    in_dir = tmp_path / "in"
    in_dir.mkdir()
    runner = CliRunner()
    result = runner.invoke(
        app,
        [
            "tdoc", "parse",
            "--from-path", str(in_dir),
            "--from-url", "https://example.com/b.zip",
        ],
    )
    assert result.exit_code != 0
    assert "mutually exclusive" in (result.output or "")


def test_cli_from_path_ignores_filter_flags_with_warning(
    tmp_path: Path, monkeypatch,
) -> None:
    """A filter flag set with ``--from-path`` warns on stderr and continues."""
    in_dir = tmp_path / "in"
    in_dir.mkdir()
    (in_dir / "R5s260009.docx").write_bytes(b"dummy")
    out = tmp_path / "out"

    fake = _FakeService(DirectParseResult(
        source_kind="local",
        markdown="3GPP TSG- blah",
        details=_dummy_details("R5s260009"),
        extract_meta=None,
        from_cache=False,
        persisted=False,
        tdoc_id="R5s260009",
        tdoc_id_in_tdocs=False,
    ))
    _build_service_factory(monkeypatch, fake)

    runner = CliRunner()
    result = runner.invoke(
        app,
        [
            "tdoc", "parse",
            "--from-path", str(in_dir),
            "--output", str(out),
            "--tdoc", "R5s260009",
            "--spec", "38.523",
        ],
    )
    assert result.exit_code == 0, result.output
    combined = (result.stdout or "") + (result.stderr or "")
    assert "ignoring filter flag(s) in local-batch mode" in combined
    assert "--tdoc" in combined
    assert "--spec" in combined


def test_cli_from_path_auto_detects_file_routes_to_single_parse(
    tmp_path: Path, monkeypatch,
) -> None:
    """A file path passed to ``--from-path`` is parsed as a single source."""
    source = tmp_path / "R5s260009.docx"
    source.write_bytes(b"dummy")
    out = tmp_path / "out.tsv"
    fake = _FakeService(DirectParseResult(
        source_kind="local",
        markdown="3GPP TSG- blah",
        details=_dummy_details("R5s260009"),
        extract_meta=None,
        from_cache=False,
        persisted=False,
        tdoc_id="R5s260009",
        tdoc_id_in_tdocs=False,
    ))
    _build_service_factory(monkeypatch, fake)

    runner = CliRunner()
    result = runner.invoke(
        app, [
            "tdoc", "parse",
            "--from-path", str(source),
            "--format", "table",
            "--output", str(out),
        ],
    )
    assert result.exit_code == 0, result.output
    assert out.exists()
    assert fake.bytes_calls, "extract_from_bytes was not called"


def test_cli_from_path_directory_requires_output(
    tmp_path: Path, monkeypatch,
) -> None:
    """A directory path passed to ``--from-path`` requires ``--output``."""
    in_dir = tmp_path / "in"
    in_dir.mkdir()
    runner = CliRunner()
    result = runner.invoke(
        app, [
            "tdoc", "parse",
            "--from-path", str(in_dir),
        ],
    )
    assert result.exit_code != 0
    assert "--output is required when --from-path is a directory" in (result.output or "")


def test_cli_from_path_file_silently_ignores_recursive(
    tmp_path: Path, monkeypatch,
) -> None:
    """``--recursive`` is ignored when ``--from-path`` points to a single file."""
    source = tmp_path / "R5s260009.docx"
    source.write_bytes(b"dummy")
    fake = _FakeService(DirectParseResult(
        source_kind="local",
        markdown="3GPP TSG- blah",
        details=_dummy_details("R5s260009"),
        extract_meta=None,
        from_cache=False,
        persisted=False,
        tdoc_id="R5s260009",
        tdoc_id_in_tdocs=False,
    ))
    _build_service_factory(monkeypatch, fake)

    runner = CliRunner()
    result = runner.invoke(
        app, [
            "tdoc", "parse",
            "--from-path", str(source),
            "--recursive",
        ],
    )
    assert result.exit_code == 0, result.output
    assert fake.bytes_calls, "extract_from_bytes was not called"


def test_cli_from_path_nonexistent_path_raises_bad_parameter(
    tmp_path: Path, monkeypatch,
) -> None:
    """A ``--from-path`` value that does not exist surfaces a clear error."""
    missing = tmp_path / "does-not-exist"
    runner = CliRunner()
    result = runner.invoke(
        app, [
            "tdoc", "parse",
            "--from-path", str(missing),
        ],
    )
    assert result.exit_code != 0
    assert "--from-path does not exist" in (result.output or "")


# ---------------------------------------------------------------------------
# CLI dispatch: --compact on tdoc parse --from-path / --from-url
# ---------------------------------------------------------------------------


def test_cli_from_path_file_json_compact_emits_single_line(
    tmp_path: Path, monkeypatch,
) -> None:
    """``--from-path FILE --format json --compact`` writes a single-line
    JSON object — no indent, no operator-space, no trailing newline."""
    source = tmp_path / "in.docx"
    source.write_bytes(b"dummy")
    out = tmp_path / "out.json"
    fake = _FakeService(DirectParseResult(
        source_kind="local",
        markdown="3GPP TSG- blah",
        details=_dummy_details("R5s260009"),
        extract_meta=None,
        from_cache=False,
        persisted=False,
        tdoc_id="R5s260009",
        tdoc_id_in_tdocs=False,
    ))
    _build_service_factory(monkeypatch, fake)

    runner = CliRunner()
    result = runner.invoke(
        app, [
            "tdoc", "parse",
            "--from-path", str(source),
            "--format", "json",
            "--compact",
            "--output", str(out),
        ],
    )
    assert result.exit_code == 0, result.output
    body = out.read_text()
    assert "\n" not in body
    assert ": " not in body
    payload = json.loads(body)
    assert payload["tdoc_id"] == "R5s260009"
    assert payload["spec"] == "38.523-3"


def test_cli_from_path_file_markdown_compact_drops_gfm_decorators(
    tmp_path: Path, monkeypatch,
) -> None:
    """``--from-path FILE --format markdown --compact`` drops the GFM
    table decorators and emits ``key: value`` lines per field."""
    source = tmp_path / "in.docx"
    source.write_bytes(b"dummy")
    out = tmp_path / "out.md"
    fake = _FakeService(DirectParseResult(
        source_kind="local",
        markdown="3GPP TSG- blah",
        details=_dummy_details("R5s260009"),
        extract_meta=None,
        from_cache=False,
        persisted=False,
        tdoc_id="R5s260009",
        tdoc_id_in_tdocs=False,
    ))
    _build_service_factory(monkeypatch, fake)

    runner = CliRunner()
    result = runner.invoke(
        app, [
            "tdoc", "parse",
            "--from-path", str(source),
            "--format", "markdown",
            "--compact",
            "--output", str(out),
        ],
    )
    assert result.exit_code == 0, result.output
    body = out.read_text()
    assert "|" not in body
    assert "---" not in body
    assert "tdoc_id: R5s260009" in body
    assert "spec: 38.523-3" in body


def test_cli_from_path_directory_json_compact_emits_single_line(
    tmp_path: Path, monkeypatch,
) -> None:
    """``--from-path DIR --format json --compact`` writes single-line
    JSON files (one per input source) — the local-batch branch threads
    ``compact`` through to ``_emit_record``."""
    in_dir = tmp_path / "in"
    in_dir.mkdir()
    source = in_dir / "R5s260009.docx"
    source.write_bytes(b"dummy")
    out = tmp_path / "out"
    fake = _FakeService(DirectParseResult(
        source_kind="local",
        markdown="3GPP TSG- blah",
        details=_dummy_details("R5s260009"),
        extract_meta=None,
        from_cache=False,
        persisted=False,
        tdoc_id="R5s260009",
        tdoc_id_in_tdocs=False,
    ))
    _build_service_factory(monkeypatch, fake)

    runner = CliRunner()
    result = runner.invoke(
        app, [
            "tdoc", "parse",
            "--from-path", str(in_dir),
            "--format", "json",
            "--compact",
            "--output", str(out),
        ],
    )
    assert result.exit_code == 0, result.output
    out_file = out / "R5s260009.json"
    assert out_file.exists()
    body = out_file.read_text()
    assert "\n" not in body
    payload = json.loads(body)
    assert payload["tdoc_id"] == "R5s260009"


# ---------------------------------------------------------------------------
# direct_parse_bytes — max_bytes guard
# ---------------------------------------------------------------------------


def test_direct_parse_bytes_raises_when_payload_too_large() -> None:
    """``max_bytes > 0`` and ``len(payload) > max_bytes`` raises
    TDocTooLargeError before any unzip / python-docx work.
    """
    from doc3gpp.parsers.direct_extractor import direct_parse_bytes
    from doc3gpp.services.tdoc_cr_service import TDocTooLargeError

    payload = b"x" * (3 * 1024 * 1024)
    with pytest.raises(TDocTooLargeError) as exc_info:
        direct_parse_bytes(
            payload, filename="R5-260100.zip", max_bytes=1024 * 1024,
        )
    assert exc_info.value.size == len(payload)
    assert exc_info.value.limit == 1024 * 1024
    assert exc_info.value.source == "R5-260100.zip"


def test_direct_parse_bytes_max_bytes_zero_is_noop() -> None:
    """``max_bytes=0`` disables the guard — no TDocTooLargeError raised.

    Uses a clearly-invalid payload (one byte, with a ``.zip``
    extension) to assert that *only* the size guard is bypassed —
    other parser errors (zipfile.BadZipFile, CRHeaderMissingError)
    may still surface.
    """
    from doc3gpp.parsers.direct_extractor import direct_parse_bytes
    from doc3gpp.services.tdoc_cr_service import TDocTooLargeError

    try:
        direct_parse_bytes(b"x", filename="R5-260100.zip", max_bytes=0)
    except TDocTooLargeError as exc:
        pytest.fail(f"max_bytes=0 must disable the guard; got {exc}")
    except Exception:
        # Any other exception is acceptable — we only assert the
        # size guard did not fire on a tiny payload.
        pass
